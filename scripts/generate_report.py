"""生成国内 viaim 迭代测试报告 .docx

以 630版本/测试报告_国内viaim_V6.3.0.docx 作为样式模板，结合：
- 缺陷 Excel（调用 parse_bugs 模块提取统计）
- 简约版测试点 .md（解析为测试范围表）
- 版本号 / 测试周期 / 参与人员 / 风险描述

原位替换模板中的动态内容，产出新一版 .docx 报告。
静态样式（字体、表格边框、页眉页脚、封面、公司信息等）完全保留。

用法示例：
    python scripts/generate_report.py \
      --version 6.4.0 \
      --cycle "2026.4.14 - 2026.5.10" \
      --excel  "640版本/iFLYBUDS--2026-05-10.xlsx" \
      --points "640版本/640简约版测试点.md" \
      --risk   "主要为交互不佳改动范围影响较大的问题" \
      --pm "李倩" --design "唐淑娴" \
      --dev "徐康,郑文祥,孙学林,万宗咏,孙庆,吴胜,廖金龙" \
      --qa  "万永琪,罗娜,马鑫鑫" \
      --out "640版本/测试报告_国内viaim_V6.4.0.docx"
"""

import argparse
import os
import re
import sys
from copy import deepcopy
from datetime import date
from pathlib import Path

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from parse_bugs import load_bugs  # noqa: E402

from docx import Document  # noqa: E402

DEFAULT_TEMPLATE = '630版本/测试报告_国内viaim_V6.3.0.docx'

SEVERITIES = ['致命', '严重', '一般', '轻微']
STATUSES   = ['暂不修复', '待修复', '已闭环']
PRIORITIES = ['紧急', '高', '中', '低']
PLATFORMS  = ['Android', 'iOS', '服务端']


def parse_points(md_path):
    """解析 `xxx简约版测试点.md`，返回 [(功能点名, 覆盖要点文本)]。

    约定：
      - 一级模块使用 `## 序号. 名称` 形式
      - 覆盖要点是紧跟的无序列表（`-` 开头）
    """
    text = Path(md_path).read_text(encoding='utf-8')
    points = []
    current_name = None
    current_lines = []
    for line in text.splitlines():
        m = re.match(r'^##\s+\d+\.\s*(.+?)\s*$', line)
        if m:
            if current_name is not None:
                points.append((current_name, '\n'.join(current_lines).strip()))
            current_name = m.group(1).strip()
            current_lines = []
        elif current_name is not None and line.strip():
            if line.lstrip().startswith(('-', '*')):
                current_lines.append(line.rstrip())
    if current_name is not None:
        points.append((current_name, '\n'.join(current_lines).strip()))
    return points


def compute_stats(bugs):
    total = len(bugs)
    plat_count = {p: 0 for p in PLATFORMS + ['未知']}
    sev_stat   = {s: {st: 0 for st in STATUSES} for s in SEVERITIES}
    plat_pri   = {p: {pr: 0 for pr in PRIORITIES} for p in PLATFORMS}
    for b in bugs:
        if b['platform'] in plat_count:
            plat_count[b['platform']] += 1
        if b['severity'] in sev_stat:
            sev_stat[b['severity']][b['status']] += 1
        if b['platform'] in plat_pri and b['priority'] in plat_pri[b['platform']]:
            plat_pri[b['platform']][b['priority']] += 1
    not_fix = [b for b in bugs if b['status'] == '暂不修复']
    return {
        'total': total,
        'plat_count': plat_count,
        'sev_stat': sev_stat,
        'plat_pri': plat_pri,
        'not_fix': not_fix,
    }


def _set_para_text(paragraph, new_text):
    """保留段落首 run 格式，覆盖整段文本。"""
    if paragraph.runs:
        first_run = paragraph.runs[0]
        for r in paragraph.runs[1:]:
            r._element.getparent().remove(r._element)
        first_run.text = new_text
    else:
        paragraph.add_run(new_text)


def _set_cell_text(cell, new_text):
    """清空单元格并填入新文本，保留首段首 run 的字体样式。多行文本用换行分隔。"""
    paragraphs = list(cell.paragraphs)
    if not paragraphs:
        cell.add_paragraph(new_text)
        return

    p0 = paragraphs[0]
    for extra in paragraphs[1:]:
        extra._element.getparent().remove(extra._element)

    # 捕获首个 run 的 rPr（字体样式）用于复制
    ref_rPr = None
    if p0.runs:
        ref_rPr = deepcopy(p0.runs[0]._element.find(
            '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr'
        ))
    for r in list(p0.runs):
        r._element.getparent().remove(r._element)

    lines = (new_text or '').split('\n')
    for i, line in enumerate(lines):
        if i > 0:
            br_run = p0.add_run()
            br_run.add_break()
        run = p0.add_run(line)
        if ref_rPr is not None:
            new_rPr = deepcopy(ref_rPr)
            existing = run._element.find(
                '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr'
            )
            if existing is not None:
                run._element.remove(existing)
            run._element.insert(0, new_rPr)


def _clone_row(table, template_row):
    """基于模板行 deepcopy 一个新行追加到表尾，返回新行对象。"""
    new_tr = deepcopy(template_row._tr)
    table._tbl.append(new_tr)
    return table.rows[-1]


def _reset_data_rows(table, keep_header=1):
    """保留表头和首条数据行，删除其余数据行。
    返回保留下来的首条数据行（作为克隆模板），若没有数据行返回 None。
    """
    data_rows = list(table.rows)[keep_header:]
    if not data_rows:
        return None
    for r in data_rows[1:]:
        r._tr.getparent().remove(r._tr)
    return list(table.rows)[keep_header]


def fill_version_info(doc, version, cycle, year_month):
    for para in doc.paragraphs:
        txt = para.text.strip()
        if re.match(r'^V\d+\.\d+\.\d+$', txt):
            _set_para_text(para, f'V{version}')
        elif re.match(r'^\d{4}年\d+月$', txt):
            _set_para_text(para, year_month)
        elif txt.startswith('版本号：'):
            _set_para_text(para, f'版本号：V{version}')
        elif txt.startswith('测试周期：'):
            _set_para_text(para, f'测试周期：{cycle}')


def fill_participants(doc, pm, design, devs, qas):
    mapping = [
        ('产品：', f'产品：{pm}'),
        ('视觉：', f'视觉：{design}'),
        ('研发：', f'研发：{"、".join(devs)}'),
        ('测试：', f'测试：{"、".join(qas)}'),
    ]
    for para in doc.paragraphs:
        txt = para.text.strip()
        for prefix, new_val in mapping:
            if txt.startswith(prefix):
                _set_para_text(para, new_val)
                break


def fill_overview_and_risk(doc, stats, risk_reason):
    total = stats['total']
    a = stats['plat_count']['Android']
    i = stats['plat_count']['iOS']
    s = stats['plat_count']['服务端']
    nf = len(stats['not_fix'])
    plat_line_new = f'Android端{a}个，iOS端{i}个，服务端{s}个'

    for para in doc.paragraphs:
        txt = para.text.strip()
        if re.match(r'^本次测试总缺陷数量\d+个$', txt):
            _set_para_text(para, f'本次测试总缺陷数量{total}个')
        elif re.match(r'^Android端\d+个，iOS端\d+个，服务端\d+个$', txt):
            _set_para_text(para, plat_line_new)
        elif txt.startswith('遗留') and '个暂不修复问题' in txt:
            _set_para_text(para, f'遗留{nf}个暂不修复问题，非【严重】缺陷，{risk_reason}')
        elif re.match(r'^总缺陷\d+个，', txt):
            _set_para_text(para, f'总缺陷{total}个，其中"暂不修复"缺陷无【严重】缺陷：')


def append_version_record(doc, version, author, change_reason):
    t = doc.tables[0]
    last_row = t.rows[-1]
    last_seq = last_row.cells[0].text.strip()
    try:
        next_seq = f'{int(float(last_seq)) + 1}.0'
    except ValueError:
        next_seq = f'{len(t.rows)}.0'

    today = date.today()
    today_str = f'{today.year}年{today.month}月{today.day}日'

    new_row = _clone_row(t, last_row)
    cells = new_row.cells
    _set_cell_text(cells[0], next_seq)
    _set_cell_text(cells[1], today_str)
    _set_cell_text(cells[2], author)
    _set_cell_text(cells[3], change_reason)
    _set_cell_text(cells[4], '')


def fill_test_scope_table(doc, points):
    t = doc.tables[2]
    template_row = _reset_data_rows(t, keep_header=1)
    if template_row is None or not points:
        return
    for _ in range(len(points) - 1):
        _clone_row(t, template_row)
    data_rows = list(t.rows)[1:]
    for idx, ((name, coverage), row) in enumerate(zip(points, data_rows), 1):
        cells = row.cells
        _set_cell_text(cells[0], str(idx))
        _set_cell_text(cells[1], name)
        _set_cell_text(cells[2], coverage)
        _set_cell_text(cells[3], '通过')


def fill_severity_table(doc, stats):
    t = doc.tables[3]
    sev_stat = stats['sev_stat']
    labels = ['致命缺陷', '严重缺陷', '一般缺陷', '轻微缺陷']
    for i, (sev, label) in enumerate(zip(SEVERITIES, labels)):
        row = t.rows[1 + i]
        vals = [sev_stat[sev][st] for st in STATUSES]
        _set_cell_text(row.cells[0], label)
        for j, v in enumerate(vals):
            _set_cell_text(row.cells[1 + j], str(v))
        _set_cell_text(row.cells[4], str(sum(vals)))

    total_row = t.rows[5]
    col_totals = [sum(sev_stat[sev][st] for sev in SEVERITIES) for st in STATUSES]
    _set_cell_text(total_row.cells[0], '总计')
    for j, v in enumerate(col_totals):
        _set_cell_text(total_row.cells[1 + j], str(v))
    _set_cell_text(total_row.cells[4], str(sum(col_totals)))


def fill_platform_priority_table(doc, stats):
    t = doc.tables[4]
    total = stats['total']
    plat_pri = stats['plat_pri']
    labels = ['Android端', 'iOS端', '服务端']
    for i, (plat, label) in enumerate(zip(PLATFORMS, labels)):
        row = t.rows[1 + i]
        vals = [plat_pri[plat][p] for p in PRIORITIES]
        subtotal = sum(vals)
        pct = f'{subtotal / total * 100:.1f}%' if total else '0.0%'
        _set_cell_text(row.cells[0], label)
        for j, v in enumerate(vals):
            _set_cell_text(row.cells[1 + j], str(v))
        _set_cell_text(row.cells[5], pct)


def fill_known_issues_table(doc, stats):
    t = doc.tables[5]
    template_row = _reset_data_rows(t, keep_header=1)
    if template_row is None:
        return
    not_fix = stats['not_fix']
    if not not_fix:
        _set_cell_text(template_row.cells[0], '1')
        _set_cell_text(template_row.cells[1], '无')
        for j in range(2, 6):
            _set_cell_text(template_row.cells[j], '/')
        return

    for _ in range(len(not_fix) - 1):
        _clone_row(t, template_row)
    data_rows = list(t.rows)[1:]
    for idx, (bug, row) in enumerate(zip(not_fix, data_rows), 1):
        cells = row.cells
        _set_cell_text(cells[0], str(idx))
        _set_cell_text(cells[1], bug['title'])
        _set_cell_text(cells[2], bug['severity'])
        _set_cell_text(cells[3], '')
        _set_cell_text(cells[4], '暂不处理')
        _set_cell_text(cells[5], '/')


def main():
    ap = argparse.ArgumentParser(description='生成国内viaim迭代测试报告 .docx')
    ap.add_argument('--version', required=True, help='版本号，如 6.4.0（不含 V 前缀）')
    ap.add_argument('--cycle', required=True, help='测试周期，如 "2026.4.14 - 2026.5.10"')
    ap.add_argument('--excel', required=True, help='缺陷 Excel 路径')
    ap.add_argument('--points', required=True, help='xxx简约版测试点.md 路径')
    ap.add_argument('--out', required=True, help='输出 .docx 路径')
    ap.add_argument('--risk', required=True, help='风险评估原因描述（嵌入风险评估段落）')
    ap.add_argument('--pm', required=True, help='产品负责人姓名')
    ap.add_argument('--design', required=True, help='视觉负责人姓名')
    ap.add_argument('--dev', required=True, help='研发成员，逗号分隔')
    ap.add_argument('--qa', required=True, help='测试成员，逗号分隔')
    ap.add_argument('--year-month', default=None, help='封面年月，默认取当前月')
    ap.add_argument('--template', default=DEFAULT_TEMPLATE, help='模板 .docx 路径（默认 V6.3.0 报告）')
    ap.add_argument('--version-author', default='万永琪', help='版本记录表记录人（默认 万永琪）')
    ap.add_argument('--change-reason', default=None, help='版本记录变更原因（默认 V{version}版本）')
    args = ap.parse_args()

    for label, path in [('模板', args.template), ('缺陷 Excel', args.excel), ('简约版测试点', args.points)]:
        if not os.path.isfile(path):
            print(f'错误：{label}文件不存在 {path}', file=sys.stderr)
            sys.exit(1)

    today = date.today()
    year_month = args.year_month or f'{today.year}年{today.month}月'
    change_reason = args.change_reason or f'V{args.version}版本'

    bugs = load_bugs(args.excel)
    stats = compute_stats(bugs)
    points = parse_points(args.points)

    doc = Document(args.template)
    fill_version_info(doc, args.version, args.cycle, year_month)
    fill_participants(
        doc,
        pm=args.pm,
        design=args.design,
        devs=[x.strip() for x in args.dev.split(',') if x.strip()],
        qas=[x.strip() for x in args.qa.split(',') if x.strip()],
    )
    fill_overview_and_risk(doc, stats, args.risk)
    append_version_record(doc, args.version, args.version_author, change_reason)
    fill_test_scope_table(doc, points)
    fill_severity_table(doc, stats)
    fill_platform_priority_table(doc, stats)
    fill_known_issues_table(doc, stats)

    out_dir = os.path.dirname(args.out)
    if out_dir and not os.path.isdir(out_dir):
        os.makedirs(out_dir, exist_ok=True)
    doc.save(args.out)

    print(f'报告已生成：{args.out}')
    print(f'  总缺陷数：{stats["total"]}')
    print(f'  Android: {stats["plat_count"]["Android"]}  '
          f'iOS: {stats["plat_count"]["iOS"]}  '
          f'服务端: {stats["plat_count"]["服务端"]}')
    print(f'  暂不修复：{len(stats["not_fix"])}')
    print(f'  功能点：{len(points)}')


if __name__ == '__main__':
    main()
